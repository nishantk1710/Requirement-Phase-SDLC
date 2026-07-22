"""Export the generated handoff Markdown to Word (.docx).

A lightweight, self-contained Markdown -> .docx converter tuned to the documents this system emits
(SRS / RTM / open-questions / seed-models): headings, paragraphs, bullet & numbered lists, tables,
code fences, blockquotes, and inline **bold** / `code`. Mermaid fenced blocks render as monospace
text (the diagram itself stays in the .md).

BEST-EFFORT BY DESIGN: if python-docx is unavailable, or any single document fails to convert, the
.docx step is skipped silently — the .md files (the primary output) are never affected.
"""

from __future__ import annotations

import re
from pathlib import Path

# bold BEFORE single-* italic so "**x**" isn't mis-split; italic excludes inner '*'/newline.
_INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
_TABLE_SEP = re.compile(r"^\s*\|[\s:|-]+\|\s*$")


def _add_runs(paragraph, text: str) -> None:
    """Render inline markdown (**bold**, *italic*, `code`, [links]) into a paragraph's runs.
    `\\|` (an escaped table pipe) is unescaped back to a literal '|'."""
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            paragraph.add_run(tok[2:-2].replace("\\|", "|")).bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            paragraph.add_run(tok[1:-1].replace("\\|", "|")).italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1].replace("\\|", "|"))
            r.font.name = "Consolas"
        elif tok.startswith("[") and "](" in tok:
            paragraph.add_run(tok[1:tok.index("]")])  # link label only (docx has no live target here)
        else:
            paragraph.add_run(tok.replace("\\|", "|"))


def _flush_table(doc, rows: list[str], spec: dict | None = None) -> None:
    from docx.shared import Pt

    from .docx_style import DOCX_STYLE, shade_header_cell

    spec = spec or DOCX_STYLE
    tcfg = spec.get("table", {})
    name = spec.get("font", "Times New Roman")
    rows = [r for r in rows if not _TABLE_SEP.match(r)]  # drop the |---|---| separator row
    if not rows:
        return
    # split on UNescaped pipes only — an escaped '\|' inside a cell is real content, not a column
    # boundary (otherwise it adds a phantom column and shifts the whole table). _add_runs unescapes.
    cells = [[c.strip() for c in re.split(r"(?<!\\)\|", r.strip().strip("|"))] for r in rows]
    ncol = max(len(r) for r in cells)
    table = doc.add_table(rows=0, cols=ncol)
    try:
        table.style = "Table Grid"          # uniform borders; header styled below (Part K)
    except Exception:
        pass
    for i, row in enumerate(cells):
        row = row + [""] * (ncol - len(row))
        tcells = table.add_row().cells
        for j, val in enumerate(row):
            para = tcells[j].paragraphs[0]
            _add_runs(para, val)
            if i == 0:                        # header row — white text on slate fill
                shade_header_cell(tcells[j], fill=tcfg.get("header_fill", "1F3A5F"),
                                  text=tcfg.get("header_text", "FFFFFF"),
                                  size=tcfg.get("header_size", 10.5), name=name)
            else:                             # body cell — Times New Roman at cell size
                for run in para.runs:
                    run.font.name = name
                    run.font.size = Pt(tcfg.get("cell_size", 10.5))


def _render_title_page(doc, lines: list[str], spec: dict | None = None) -> None:
    """Render the SRS title block as a CENTRED, page-1 title page, then a page break. Lines come from
    the <!--TITLEPAGE--> marker block: the '# …' line is the document title; a lone 'for' labels the
    project name on the next line; a 'Document format …' line is the small disclaimer; the rest are
    centred meta lines (version, prepared-by, date). Uses direct run formatting (not Heading styles)
    so nothing here appears in the Table of Contents."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    from .docx_style import DOCX_STYLE

    spec = spec or DOCX_STYLE
    name = spec.get("font", "Times New Roman")
    styles = spec.get("styles", {})
    title_hex = styles.get("Title", {}).get("color", "1F3A5F")
    accent_hex = styles.get("Heading 1", {}).get("color", title_hex)
    muted_hex, body_hex = "6B6B6B", "1A1A1A"

    lines = [ln for ln in lines if ln.strip()]
    for_idx = next((i for i, ln in enumerate(lines) if ln.strip().lower() == "for"), None)
    project_idx = for_idx + 1 if for_idx is not None else None

    for i, ln in enumerate(lines):
        text = ln.strip()
        is_title = text.startswith("#")
        if is_title:
            text = text.lstrip("#").strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = name
        if i == 0:
            p.paragraph_format.space_before = Pt(140)          # push the block toward the page centre
        if is_title:
            run.bold = True
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor.from_string(title_hex)
            p.paragraph_format.space_after = Pt(4)
        elif i == for_idx:
            run.italic = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor.from_string(muted_hex)
        elif i == project_idx:
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor.from_string(accent_hex)
            p.paragraph_format.space_after = Pt(12)
        elif text.startswith("Document format"):
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(muted_hex)
            p.paragraph_format.space_before = Pt(28)
        else:
            run.font.size = Pt(12.5)
            run.font.color.rgb = RGBColor.from_string(body_hex)
    doc.add_page_break()


def _insert_toc_field(doc) -> None:
    """Insert a live Word Table-of-Contents FIELD (headings 1-3, hyperlinked, with page numbers).
    Word fills it in — with page numbers — when fields update (see `_enable_field_update`)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = doc.add_paragraph().add_run()
    r = run._r
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = 'Table of Contents — right-click and choose "Update Field" (or press F9) to fill in the page numbers.'
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, placeholder, end):
        r.append(el)


def _enable_field_update(doc) -> None:
    """Ask Word to refresh fields (TOC page numbers, footer PAGE) when the document is opened."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields"); el.set(qn("w:val"), "true")
        settings.append(el)


def markdown_to_docx(md_text: str, out_path: str | Path) -> None:
    """Convert a Markdown string to a STYLED .docx (Part K named styles: Times New Roman, size/colour
    hierarchy, table header fill, page numbers). Raises if python-docx is missing."""
    import docx
    from docx.shared import Pt

    from .docx_style import apply_document_styles

    doc = docx.Document()
    spec = apply_document_styles(doc)          # Times New Roman + size/colour hierarchy + footer

    in_code, code_buf, table_buf = False, [], []
    in_title, title_buf = False, []       # <!--TITLEPAGE-->…<!--/TITLEPAGE-->
    in_toc, toc_inserted = False, False    # <!--TOC-->…<!--/TOC--> (replaced by a live TOC field)
    for raw in (md_text or "").splitlines():
        line = raw.rstrip("\n")
        s0 = line.strip()

        # --- title page block: collect, then render centred + page break ------------------------
        if s0 == "[[TITLEPAGE]]":
            if table_buf:
                _flush_table(doc, table_buf, spec); table_buf = []
            in_title, title_buf = True, []
            continue
        if s0 == "[[/TITLEPAGE]]":
            in_title = False
            _render_title_page(doc, title_buf, spec)
            title_buf = []
            continue
        if in_title:
            title_buf.append(line)
            continue

        # --- table of contents: a live Word field replaces the .md outline ----------------------
        if s0 == "[[TOC]]":
            if table_buf:
                _flush_table(doc, table_buf, spec); table_buf = []
            _insert_toc_field(doc); toc_inserted = True
            in_toc = True
            continue
        if s0 == "[[/TOC]]":
            in_toc = False
            doc.add_page_break()
            continue
        if in_toc:
            continue                       # skip the .md outline lines — the field renders instead

        if line.strip().startswith("```"):
            if in_code:
                run = doc.add_paragraph().add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                code_buf = []
            in_code = not in_code
            continue
        if in_code:
            code_buf.append(line)
            continue

        if line.strip().startswith("|"):
            table_buf.append(line)
            continue
        if table_buf:
            _flush_table(doc, table_buf, spec)
            table_buf = []

        s = line.strip()
        if not s or s == "---":
            continue
        if s.startswith("#"):
            level = len(s) - len(s.lstrip("#"))
            _add_runs(doc.add_heading("", level=0 if level == 1 else min(level - 1, 4)), s.lstrip("#").strip())
            continue
        if s.startswith("> "):
            p = doc.add_paragraph()
            try:
                p.style = "Intense Quote"
            except Exception:
                pass
            _add_runs(p, s[2:])
            continue
        m = re.match(r"^(\d+)\.\s+(.*)", s)
        if m:
            # keep the source's OWN number as literal text (a plain indented paragraph). Word's
            # "List Number" style runs one continuous counter across the whole document, which turns
            # each block's "1./2./3." into 4/5/6, 7/8/9… — losing the intended per-block numbering (D-F6).
            para = doc.add_paragraph()
            try:
                para.paragraph_format.left_indent = Pt(18)
            except Exception:
                pass
            _add_runs(para, f"{m.group(1)}. {m.group(2)}")
            continue
        if s.startswith("- ") or s.startswith("* "):
            _add_runs(doc.add_paragraph(style="List Bullet"), s[2:])
            continue
        _add_runs(doc.add_paragraph(), s)

    if table_buf:
        _flush_table(doc, table_buf, spec)
    if toc_inserted:
        _enable_field_update(doc)          # Word fills TOC page numbers on open
    doc.save(str(out_path))


def write_docx_versions(out_dir: str | Path, md_files: dict[str, str]) -> list[str]:
    """Write a `.docx` beside each `.md` entry in `md_files`. Best-effort — never raises.

    Returns the list of .docx filenames written (empty if python-docx is unavailable)."""
    out_dir = Path(out_dir)
    written: list[str] = []
    try:
        import docx  # noqa: F401  — probe once; skip cleanly if not installed
    except ImportError:
        return written
    for name, content in md_files.items():
        if not name.endswith(".md"):
            continue
        docx_name = name[:-3] + ".docx"
        try:
            markdown_to_docx(content, out_dir / docx_name)
            written.append(docx_name)
        except Exception:
            pass  # a single doc failing must not break the rest or the .md output
    return written
