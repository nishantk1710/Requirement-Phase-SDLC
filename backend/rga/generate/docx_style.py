"""Word (.docx) document styling for the SRS/RTM deliverables (Part K).

Applies NAMED Word styles (Normal, Title, Heading 1/2/3, Table Header) so the document is clean and
re-themeable — never ad-hoc run styling. Times New Roman throughout (incl. the East-Asian /
complex-script font slots so Word cannot fall back to Calibri), a consistent size hierarchy, and
defined heading/body/table colours. Page numbers in the footer.

These are DOCUMENT styles for this deliverable — deliberately independent of the product design
tokens rendered in SRS §3.1 (which describe the product's UI, not this document).
"""

from __future__ import annotations

# Default scheme (Part K). Mirrored in config.yaml under `docx:`; `apply_document_styles` accepts an
# override so it stays configurable. Colours are hex WITHOUT '#'. Sizes are points.
DOCX_STYLE: dict = {
    "font": "Times New Roman",
    "styles": {
        # word_style_name : {size, bold, color}
        "Title":     {"size": 24.0, "bold": True, "color": "1F3A5F"},   # deep navy/slate
        "Heading 1": {"size": 18.0, "bold": True, "color": "1F3A5F"},
        "Heading 2": {"size": 14.0, "bold": True, "color": "2E4A6B"},   # lighter slate
        "Heading 3": {"size": 12.5, "bold": True, "color": "3A3A3A"},   # dark grey
        "Normal":    {"size": 11.0, "bold": False, "color": "1A1A1A"},  # body near-black
    },
    "table": {"cell_size": 10.5, "header_size": 10.5,
              "header_text": "FFFFFF", "header_fill": "1F3A5F"},   # white on slate
    "caption": {"size": 10.0, "italic": True, "color": "3A3A3A"},
    "heading_space_before": 10.0, "heading_space_after": 4.0,
}


_RESOLVED_SPEC: dict | None = None


def _deep_merge(base: dict, over: dict) -> dict:
    out = dict(base)
    for k, v in (over or {}).items():
        out[k] = _deep_merge(out[k], v) if isinstance(v, dict) and isinstance(out.get(k), dict) else v
    return out


def resolved_spec() -> dict:
    """DOCX_STYLE deep-merged with any `docx:` overrides in config.yaml (best-effort, cached)."""
    global _RESOLVED_SPEC
    if _RESOLVED_SPEC is None:
        override: dict = {}
        try:
            from ..config import load_config
            override = load_config().docx or {}
        except Exception:
            override = {}
        _RESOLVED_SPEC = _deep_merge(DOCX_STYLE, override)
    return _RESOLVED_SPEC


def _set_style_font(style, *, name, size=None, bold=None, italic=None, color=None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    f = style.font
    f.name = name
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if color is not None:
        f.color.rgb = RGBColor.from_string(color)
    # Force ALL font slots so Word does not silently substitute Calibri for headings / CJK / symbols.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(slot), name)


def _add_page_number_footer(doc) -> None:
    """Centre a 'Page N' field in every section footer."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = ""
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("Page ")
        for kind, txt in (("begin", None), (None, "PAGE"), ("end", None)):
            if kind:
                fld = OxmlElement("w:fldChar")
                fld.set(qn("w:fldCharType"), kind)
                run._r.append(fld)
            else:
                instr = OxmlElement("w:instrText")
                instr.set(qn("xml:space"), "preserve")
                instr.text = txt
                run._r.append(instr)


def apply_document_styles(doc, spec: dict | None = None) -> dict:
    """Apply the named-style scheme to a python-docx Document. Returns the spec used.
    When `spec` is None, uses DOCX_STYLE merged with any `docx:` overrides in config.yaml."""
    spec = spec or resolved_spec()
    name = spec["font"]
    st = doc.styles
    for style_name, cfg in spec["styles"].items():
        try:
            _set_style_font(st[style_name], name=name, size=cfg.get("size"),
                            bold=cfg.get("bold"), color=cfg.get("color"))
        except KeyError:
            continue  # a built-in style missing in this template — skip, never crash
        # consistent heading spacing
        if style_name.startswith("Heading") or style_name == "Title":
            from docx.shared import Pt
            pf = st[style_name].paragraph_format
            pf.space_before = Pt(spec.get("heading_space_before", 10.0))
            pf.space_after = Pt(spec.get("heading_space_after", 4.0))
    _add_page_number_footer(doc)
    return spec


def shade_header_cell(cell, *, fill: str, text: str, size: float, name: str) -> None:
    """White-on-slate styling for a table header cell (fill + run colour/size), Part K."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor.from_string(text)
            run.font.size = Pt(size)
            run.font.name = name
            run.bold = True
