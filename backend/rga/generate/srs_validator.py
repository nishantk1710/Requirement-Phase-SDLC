"""SRS format/schema validator — the format-conformance gate before handoff.

The Design team parses the generated SRS with a STRICT deterministic parser (it keys off Word
heading styles, section numbers, specific tables, and `REQ-/NFR-/BR-n` tags). If the SRS drifts
from that format, their parser silently drops artifacts. This validator mirrors that contract and
checks every generated SRS against a REFERENCE SCHEMA (`srs_format_schema.json`) so a malformed
spec can never reach Design unnoticed.

It is fully DETERMINISTIC (no LLM): it walks the document exactly as the parser does — the same
heading detection, the same table/paragraph interleaving — and evaluates the schema's rules.

Works on either the Markdown (`.md`, what the generator emits) or the Word `.docx` (what Design
actually receives). Public API:
    validate_markdown(md, schema=None) -> report dict
    validate_docx(path, schema=None)   -> report dict
    validate_srs_file(path, schema=None) -> report dict   (dispatch by extension)
    load_schema(path=None) -> dict
`report["ok"]` is True iff every check passes; `report["violations"]` lists what failed.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

DEFAULT_SCHEMA_PATH = Path(__file__).with_name("srs_format_schema.json")

_SECNUM_RE = re.compile(r"^(\d+(?:\.\d+)*)")
_HEADING_STYLE_RE = re.compile(r"^Heading (\d+)$")
_TAG_RE = re.compile(r"^(REQ|NFR|BR)-(\d+)", re.IGNORECASE)
_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


# --------------------------------------------------------------------------- #
# Intermediate structure — the same view of the document the parser builds.
# --------------------------------------------------------------------------- #
class _Doc:
    def __init__(self) -> None:
        self.headings: list[dict] = []      # {num, title, level}
        self.tables: list[dict] = []        # {sec_num, sec_title, header:[...], data_rows:int}
        self.paragraphs: list[dict] = []    # {sec_num, sec_title, text}
        self.tags: dict[str, int] = {"REQ": 0, "NFR": 0, "BR": 0}

    def has_section(self, num: str = "", title_contains: str = "") -> tuple[bool, str]:
        """Return (found, title-of-match). Match by section number if given, else by title."""
        for h in self.headings:
            if num and h["num"] == num:
                return True, h["title"]
            if not num and title_contains and title_contains.lower() in h["title"].lower():
                return True, h["title"]
        return False, ""


def _split_num_title(text: str) -> tuple[str, str]:
    """'2.3 User Classes' -> ('2.3', 'User Classes'); 'Appendix A: Glossary' -> ('', 'Appendix A: Glossary')."""
    text = " ".join((text or "").split())
    m = _SECNUM_RE.match(text)
    num = m.group(1) if m else ""
    title = text[len(num):].lstrip(" .\t:") if num else text
    return num, title.strip()


def _record_tag(doc: _Doc, text: str) -> None:
    m = _TAG_RE.match(text.strip())
    if m:
        doc.tags[m.group(1).upper()] += 1


# --------------------------------------------------------------------------- #
# Markdown extraction (what the generator emits)
# --------------------------------------------------------------------------- #
def _cells(line: str) -> list[str]:
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


def _is_pipe_row(line: str) -> bool:
    return line.strip().startswith("|")


def _is_separator(line: str) -> bool:
    s = line.strip()
    return bool(s) and set(s) <= set("|-: ") and "-" in s


def extract_from_markdown(md: str) -> _Doc:
    doc = _Doc()
    lines = md.splitlines()
    cur_num, cur_title = "", ""
    i = 0
    while i < len(lines):
        line = lines[i]
        # table: a pipe row immediately followed by a separator row
        if _is_pipe_row(line) and i + 1 < len(lines) and _is_separator(lines[i + 1]):
            header = _cells(line)
            data = 0
            j = i + 2
            while j < len(lines) and _is_pipe_row(lines[j]) and not _is_separator(lines[j]):
                if any(_cells(lines[j])):
                    data += 1
                j += 1
            doc.tables.append({"sec_num": cur_num, "sec_title": cur_title, "header": header, "data_rows": data})
            i = j
            continue
        m = _MD_HEADING_RE.match(line)
        if m:
            num, title = _split_num_title(m.group(2))
            cur_num, cur_title = num, title
            doc.headings.append({"num": num, "title": title, "level": len(m.group(1))})
            i += 1
            continue
        text = line.strip()
        if text:
            # strip list markers + bold so a '- **REQ-1:** …' line reads as the parser sees it
            plain = re.sub(r"^[-*]\s+", "", text).replace("**", "").strip()
            _record_tag(doc, plain)
            doc.paragraphs.append({"sec_num": cur_num, "sec_title": cur_title, "text": plain})
        i += 1
    return doc


# --------------------------------------------------------------------------- #
# DOCX extraction (what Design receives) — same block walk as the parser
# --------------------------------------------------------------------------- #
def extract_from_docx(path: str | Path) -> _Doc:
    from docx import Document                      # lazy: only needed for .docx validation
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(str(path))
    doc = _Doc()
    cur_num, cur_title = "", ""
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            par = Paragraph(child, document)
            style = par.style.name if par.style else ""
            hm = _HEADING_STYLE_RE.match(style or "")
            text = par.text.strip()
            if hm:
                num, title = _split_num_title(text)
                cur_num, cur_title = num, title
                doc.headings.append({"num": num, "title": title, "level": int(hm.group(1))})
            elif text:
                _record_tag(doc, text)
                doc.paragraphs.append({"sec_num": cur_num, "sec_title": cur_title, "text": text})
        elif child.tag == qn("w:tbl"):
            tbl = Table(child, document)
            rows = tbl.rows
            header = [c.text.strip() for c in rows[0].cells] if rows else []
            data = sum(1 for r in rows[1:] if any(c.text.strip() for c in r.cells))
            doc.tables.append({"sec_num": cur_num, "sec_title": cur_title, "header": header, "data_rows": data})
    return doc


# --------------------------------------------------------------------------- #
# Schema loading + validation
# --------------------------------------------------------------------------- #
def load_schema(path: str | Path | None = None) -> dict:
    return json.loads(Path(path or DEFAULT_SCHEMA_PATH).read_text(encoding="utf-8"))


def _find_table(doc: _Doc, spec: dict) -> dict | None:
    want_num = spec.get("in_section")
    want_title = (spec.get("in_section_title_contains") or "").lower()
    for t in doc.tables:
        if want_num and t["sec_num"] == want_num:
            return t
        if want_title and want_title in t["sec_title"].lower():
            return t
    return None


def _header_ok(header: list[str], expected: list[str], mode: str) -> bool:
    hl = [h.lower().strip() for h in header]
    el = [e.lower().strip() for e in expected]
    if mode == "first_cell":
        return bool(hl) and hl[0].startswith(el[0])
    # "exact": each expected cell matches the header cell in the same position
    return len(hl) >= len(el) and all(hl[k] == el[k] for k in range(len(el)))


def validate(doc: _Doc, schema: dict, *, source: str = "") -> dict:
    violations: list[dict] = []
    total = 0

    def check(passed: bool, rule: str, detail: str) -> None:
        nonlocal total
        total += 1
        if not passed:
            violations.append({"rule": rule, "detail": detail})

    # 1) required sections (presence, + title when specified)
    for s in schema.get("required_sections", []):
        num = s.get("num", "")
        tc = s.get("title_contains", "")
        found, title = doc.has_section(num=num, title_contains=("" if num else tc))
        label = num or tc
        if num and found and tc and tc.lower() not in title.lower():
            check(False, "section", f"§{num} title should contain '{tc}' — found '{title}'")
        else:
            check(found, "section", f"missing required section: {label}")

    # 2) required tables (present in the right section, right header, enough rows)
    for t in schema.get("required_tables", []):
        loc = t.get("in_section") or t.get("in_section_title_contains") or t["id"]
        found = _find_table(doc, t)
        if not found:
            check(False, "table", f"missing required table '{t['id']}' in §{loc}")
            continue
        if not _header_ok(found["header"], t["header"], t.get("header_match", "exact")):
            check(False, "table", f"table '{t['id']}' (§{loc}) header {found['header']} != expected {t['header']}")
        elif found["data_rows"] < t.get("min_data_rows", 1):
            check(False, "table", f"table '{t['id']}' (§{loc}) has {found['data_rows']} data row(s), needs {t.get('min_data_rows', 1)}")
        else:
            check(True, "table", "")

    # 3) tagged requirements (REQ / NFR / BR counts)
    for r in schema.get("required_tagged_requirements", []):
        n = doc.tags.get(r["tag"].upper(), 0)
        check(n >= r.get("min_count", 0), "tagged",
              f"{r['tag']}-tagged lines: found {n}, need >= {r.get('min_count', 0)} ({r['id']})")

    # 4) required text patterns (scoped to a section when specified)
    for p in schema.get("required_patterns", []):
        rx = re.compile(p["regex"], re.IGNORECASE)
        scope = (p.get("in_section_title_contains") or "").lower()
        n = sum(1 for para in doc.paragraphs
                if (not scope or scope in para["sec_title"].lower()) and rx.search(para["text"]))
        check(n >= p.get("min_count", 1), "pattern",
              f"pattern '{p['id']}' matched {n} time(s), need >= {p.get('min_count', 1)}"
              + (f" in a section titled '{p['in_section_title_contains']}'" if scope else ""))

    ok = not violations
    passed = total - len(violations)
    summary = (f"OK — {passed}/{total} format checks passed"
               if ok else
               f"FAILED — {len(violations)} violation(s) of {total} checks: "
               + "; ".join(v["detail"] for v in violations[:6])
               + (" …" if len(violations) > 6 else ""))
    return {
        "ok": ok,
        "schema": f"{schema.get('name', 'schema')} v{schema.get('version', '?')}",
        "source": source,
        "checks_total": total,
        "checks_passed": passed,
        "violations": violations,
        "extracted": {"sections": len(doc.headings), "tables": len(doc.tables), **doc.tags},
        "summary": summary,
    }


def validate_markdown(md: str, schema: dict | None = None) -> dict:
    return validate(extract_from_markdown(md), schema or load_schema(), source="markdown")


def validate_docx(path: str | Path, schema: dict | None = None) -> dict:
    return validate(extract_from_docx(path), schema or load_schema(), source=str(path))


def validate_srs_file(path: str | Path, schema: dict | None = None) -> dict:
    p = Path(path)
    if p.suffix.lower() == ".docx":
        return validate_docx(p, schema)
    return validate_markdown(p.read_text(encoding="utf-8"), schema)


class FormatInvalid(RuntimeError):
    """Raised (only in strict mode) when the generated SRS violates the reference format schema."""
