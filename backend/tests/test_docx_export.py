"""Word (.docx) export of the generated handoff markdown — headings, tables, lists, code fences.
Best-effort: non-.md entries are skipped and a single failure never aborts the batch."""

from __future__ import annotations

import pytest

from rga.generate.docx_export import markdown_to_docx, write_docx_versions

MD = """# Software Requirements Specification

## 1. Introduction

Some **bold** text and `inline code`.

- bullet one
- bullet two

1. step one
2. step two

| ID | Requirement |
|---|---|
| REQ-1 | The system shall do X. |

```mermaid
flowchart LR
  a --> b
```
"""


def test_markdown_to_docx_produces_a_valid_document(tmp_path):
    docx = pytest.importorskip("docx")
    out = tmp_path / "SRS.docx"
    markdown_to_docx(MD, out)
    assert out.exists() and out.stat().st_size > 0
    d = docx.Document(str(out))
    heads = [p.text for p in d.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"]
    assert "Software Requirements Specification" in heads
    assert "1. Introduction" in heads
    assert len(d.tables) == 1
    assert [c.text for c in d.tables[0].rows[0].cells] == ["ID", "Requirement"]


def test_docx_table_handles_escaped_pipe_without_phantom_column(tmp_path):
    """D-F2: an escaped '\\|' inside a cell is content, not a column boundary — no phantom column,
    and it is unescaped back to a literal '|'."""
    docx = pytest.importorskip("docx")
    md = "| ID | Evidence |\n|---|---|\n| REQ-1 | choose yes\\|no |\n"
    out = tmp_path / "t.docx"
    markdown_to_docx(md, out)
    d = docx.Document(str(out))
    assert len(d.tables) == 1
    assert len(d.tables[0].rows[0].cells) == 2            # 2 columns, not 3
    assert d.tables[0].rows[1].cells[1].text == "choose yes|no"


def test_docx_renders_italics_and_literal_block_numbering(tmp_path):
    """D-F6: *italic* is supported, and per-block numbering ('1.' restarting) is preserved (no
    continuous Word counter turning the second block into 3/4)."""
    docx = pytest.importorskip("docx")
    md = "*None identified.*\n\n1. First\n2. Second\n\n1. Alpha\n2. Beta\n"
    out = tmp_path / "n.docx"
    markdown_to_docx(md, out)
    d = docx.Document(str(out))
    texts = [p.text for p in d.paragraphs]
    assert "None identified." in texts                    # italics asterisks stripped
    assert any(t.startswith("1. First") for t in texts)   # both blocks start at 1 (literal numbers)
    assert any(t.startswith("1. Alpha") for t in texts)


def test_docx_styling_times_new_roman_sizes_and_colours(tmp_path):
    """Part K: named styles report Times New Roman at the configured sizes/colours, and every
    heading text/fill colour pair meets WCAG AA."""
    docx = pytest.importorskip("docx")
    from rga.generate.color_util import contrast_ratio
    from rga.generate.docx_style import DOCX_STYLE

    out = tmp_path / "S.docx"
    markdown_to_docx(MD, out)
    d = docx.Document(str(out))
    exp = DOCX_STYLE["styles"]
    for sname in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        st = d.styles[sname]
        assert st.font.name == "Times New Roman", f"{sname} font"
        assert abs(float(st.font.size.pt) - exp[sname]["size"]) < 0.01, f"{sname} size"
        assert str(st.font.color.rgb) == exp[sname]["color"].upper(), f"{sname} colour"
    # WCAG AA: header text on its fill, and every heading colour on the white page
    t = DOCX_STYLE["table"]
    assert contrast_ratio(t["header_text"], t["header_fill"]) >= 4.5
    for key in exp:
        assert contrast_ratio(exp[key]["color"], "FFFFFF") >= 4.5, f"{key} on white"


def test_title_page_is_centred_paginated_and_not_headings(tmp_path):
    """The [[TITLEPAGE]] block renders as a centred page-1 title block (title, 'for', project name,
    version, prepared-by, date, disclaimer), followed by a page break, using direct formatting (no
    Heading/Title styles) so nothing here leaks into the Table of Contents."""
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    md = "\n".join([
        "[[TITLEPAGE]]",
        "# Software Requirements Specification",
        "for",
        "QuickBite",
        "Version 1.0 approved",
        "Prepared by RGA (Agentic Requirement Gathering & Analysis)",
        "3 March 2025",
        "Document format based on the IEEE SRS template. Copyright © 1999 by Karl E. Wiegers. "
        "Permission is granted to use, modify, and distribute this document.",
        "[[/TITLEPAGE]]",
        "",
        "## 1. Introduction",
        "Body.",
    ])
    out = tmp_path / "tp.docx"
    markdown_to_docx(md, out)
    d = docx.Document(str(out))
    texts = [p.text for p in d.paragraphs]
    assert "Software Requirements Specification" in texts
    assert "for" in texts and "QuickBite" in texts
    assert any("Prepared by RGA (Agentic Requirement Gathering & Analysis)" == t for t in texts)
    assert any("Karl E. Wiegers" in t and "distribute this document" in t for t in texts)
    tp = [p for p in d.paragraphs if p.text in ("Software Requirements Specification", "for", "QuickBite")]
    assert tp and all(p.alignment == WD_ALIGN_PARAGRAPH.CENTER for p in tp)     # centred
    assert all(not p.style.name.startswith("Heading") and p.style.name != "Title" for p in tp)  # not ToC-visible
    assert 'w:type="page"' in d.element.xml                                     # page break after title page


def test_table_of_contents_inserts_live_word_field(tmp_path):
    """[[TOC]] becomes a real Word TOC field (headings 1-3, page numbers on update); the .md outline
    lines are dropped from the .docx, and Word is asked to refresh fields on open."""
    docx = pytest.importorskip("docx")
    md = "\n".join([
        "## Table of Contents",
        "[[TOC]]",
        "- **1. Introduction**",
        "    - 1.1 Purpose",
        "[[/TOC]]",
        "",
        "## 1. Introduction",
        "Body.",
    ])
    out = tmp_path / "toc.docx"
    markdown_to_docx(md, out)
    d = docx.Document(str(out))
    body_xml = d.element.xml
    assert "fldChar" in body_xml and "instrText" in body_xml and "TOC" in body_xml   # a field, not text
    texts = [p.text for p in d.paragraphs]
    assert not any(t.strip() == "1.1 Purpose" for t in texts)                    # outline not duplicated
    assert "updateFields" in d.settings.element.xml                              # page numbers filled on open


def test_markerless_markdown_unaffected(tmp_path):
    """Backward-compat: markdown without the markers (e.g. the RTM) is converted exactly as before —
    no title-page centring, no page break, no TOC field."""
    docx = pytest.importorskip("docx")
    out = tmp_path / "rtm.docx"
    markdown_to_docx("# Requirements Traceability Matrix\n\n## Rows\n\nContent.", out)
    d = docx.Document(str(out))
    assert 'w:type="page"' not in d.element.xml
    assert "updateFields" not in d.settings.element.xml
    assert "Requirements Traceability Matrix" in [p.text for p in d.paragraphs]


def test_write_docx_versions_only_converts_md(tmp_path):
    pytest.importorskip("docx")
    written = write_docx_versions(tmp_path, {
        "SRS.md": "# SRS\n\nHello world.",
        "RTM.md": "# RTM\n\n| A | B |\n|---|---|\n| 1 | 2 |",
        "manifest.json": "{}",
    })
    assert set(written) == {"SRS.docx", "RTM.docx"}          # .md converted
    assert (tmp_path / "SRS.docx").exists() and (tmp_path / "RTM.docx").exists()
    assert not (tmp_path / "manifest.docx").exists()          # non-.md skipped
